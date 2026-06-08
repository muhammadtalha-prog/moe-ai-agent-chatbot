import os
import csv
import json
import io
from pathlib import Path
from agent.security import get_safe_path

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pandas as pd
except ImportError:
    pd = None

def read_file(path_str: str) -> str:
    """
    Reads the content of a file based on its file extension.
    Supports specialized formats (Word, Excel, PDF, CSV, JSON) and
    dynamically falls back to raw text decoding for any other format.
    """
    safe_path = get_safe_path(path_str)
    
    if not safe_path.exists():
        raise FileNotFoundError(f"File not found: {safe_path}")
        
    ext = safe_path.suffix.lower()
    
    # 1. Word Document (.docx)
    if ext == '.docx':
        if not docx:
            raise ImportError("python-docx is not installed. Unable to read DOCX files.")
        doc = docx.Document(safe_path)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)
        
    # 2. Excel Spreadsheets (.xlsx, .xls)
    elif ext in ['.xlsx', '.xls']:
        if not pd:
            raise ImportError("pandas is not installed. Unable to read Excel files.")
        df = pd.read_excel(safe_path)
        return df.to_csv(index=False)
        
    # 3. PDF Files
    elif ext == '.pdf':
        if not PdfReader:
            raise ImportError("pypdf is not installed. Unable to read PDF files.")
        reader = PdfReader(safe_path)
        text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text.append(t)
        return "\n".join(text)
        
    # 4. JSON Files
    elif ext == '.json':
        with open(safe_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return json.dumps(data, indent=2)
            
    # 5. CSV Files
    elif ext == '.csv':
        rows = []
        with open(safe_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(", ".join(row))
        return "\n".join(rows)
        
    # 6. Universal Text Fallback (supports .txt, .md, .py, .html, .xml, .yaml, .css, .js, etc.)
    else:
        try:
            with open(safe_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            raise IOError(f"Unable to read file '{path_str}' as a text format: {str(e)}")

def write_file(path_str: str, content: str) -> None:
    """
    Writes content back to a file, formatting it to match the extension.
    Supports txt, docx, xlsx, pdf, csv, json, and fallbacks.
    """
    safe_path = get_safe_path(path_str)
    # Create directory if it doesn't exist
    safe_path.parent.mkdir(parents=True, exist_ok=True)
    
    ext = safe_path.suffix.lower()
    
    # 1. Word Document (.docx)
    if ext == '.docx':
        if not docx:
            raise ImportError("python-docx is not installed. Unable to write DOCX files.")
        doc = docx.Document()
        lines = content.split('\n')
        for line in lines:
            doc.add_paragraph(line)
        doc.save(safe_path)
        
    # 2. Excel Spreadsheets (.xlsx, .xls)
    elif ext in ['.xlsx', '.xls']:
        if not pd:
            raise ImportError("pandas is not installed. Unable to write Excel files.")
        # Try to read the content as a CSV string into a DataFrame
        try:
            df = pd.read_csv(io.StringIO(content))
        except Exception:
            # If not standard CSV, write lines as single column rows
            lines = content.strip().split('\n')
            df = pd.DataFrame(lines, columns=["Content"])
        df.to_excel(safe_path, index=False)
        
    # 3. PDF Files
    elif ext == '.pdf':
        if not FPDF:
            raise ImportError("fpdf is not installed. Unable to write PDF files.")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        clean_content = content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, txt=clean_content)
        pdf.output(str(safe_path))
        
    # 4. JSON Files
    elif ext == '.json':
        try:
            data = json.loads(content)
            with open(safe_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
        except json.JSONDecodeError:
            with open(safe_path, 'w', encoding='utf-8') as f:
                f.write(content)
                
    # 5. CSV Files
    elif ext == '.csv':
        lines = content.strip().split('\n')
        with open(safe_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            for line in lines:
                if ',' in line:
                    writer.writerow([cell.strip() for cell in csv.reader([line]).__next__()])
                else:
                    writer.writerow([line])
                    
    # 6. Universal Text Fallback
    else:
        with open(safe_path, 'w', encoding='utf-8') as f:
            f.write(content)
